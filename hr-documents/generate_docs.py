"""Generate all HR documents as .docx files for Arafat Iqbal / Craftsmen IT."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

OUT = os.path.dirname(os.path.abspath(__file__))

NAVY = RGBColor(0x1A, 0x3C, 0x6E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = RGBColor(0xF0, 0xF4, 0xFB)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)

# ─── helpers ────────────────────────────────────────────────────────────────

def new_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)
    # default body font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    return doc

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'none'))
        tag.set(qn('w:sz'), kwargs.get('sz', '6'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', '1A3C6E'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def para_space(doc, pt_before=0, pt_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pt_before)
    p.paragraph_format.space_after  = Pt(pt_after)
    return p

def add_header_block(doc, ref_no, date_str):
    """Company header + ref/date line."""
    # Company name
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("CRAFTSMEN IT")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY
    # sub
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run("IT Services & Solutions  |  UAE & Saudi Arabia")
    r2.font.size = Pt(9)
    r2.font.color.rgb = DARK_GRAY
    # contact
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(2)
    r3 = p3.add_run("E: contact@craftsmenit.com   |   T: +966-59-787-9394   |   www.craftsmenit.com")
    r3.font.size = Pt(9)
    r3.font.color.rgb = DARK_GRAY
    # horizontal rule via border on paragraph
    add_hrule(doc)
    # Ref / Date
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Inches(3)
    t.columns[1].width = Inches(3)
    lc = t.cell(0, 0)
    rc = t.cell(0, 1)
    for c in (lc, rc):
        for b in ('top','left','bottom','right'):
            set_cell_border(c, **{b:'none'})
    lp = lc.paragraphs[0]
    lp.add_run(f"Ref No:  {ref_no}").font.size = Pt(10)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.add_run(f"Date:  {date_str}").font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_hrule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A3C6E')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_title(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    run.font.underline = True

def add_body_para(doc, text, bold_parts=None, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text).font.size = Pt(size)
    return p

def navy_table(doc, headers, rows):
    """Table with navy header row and alternating light rows."""
    col_count = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=col_count)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, '1A3C6E')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
    # data
    for ri, row_data in enumerate(rows):
        drow = t.rows[ri + 1]
        bg = 'F0F4FB' if ri % 2 == 0 else 'FFFFFF'
        is_total = row_data[0].startswith('TOTAL') or row_data[0].startswith('Total')
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            if is_total:
                set_cell_bg(cell, '1A3C6E')
                r = cell.paragraphs[0].add_run(val)
                r.bold = True
                r.font.color.rgb = WHITE
                r.font.size = Pt(10)
            else:
                set_cell_bg(cell, bg)
                cell.paragraphs[0].add_run(val).font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def add_signature_block(doc, show_stamp=True):
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    t = doc.add_table(rows=3, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for cell in row.cells:
            for b in ('top','left','bottom','right'):
                set_cell_border(cell, **{b:'none'})

    left_cell  = t.cell(0, 0)
    right_cell = t.cell(0, 1)

    if show_stamp:
        lp = left_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run("[OFFICIAL SEAL]")
        lr.font.size = Pt(9)
        lr.font.color.rgb = NAVY
        lr.font.bold = True

    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp.paragraph_format.space_before = Pt(30)
    rp.add_run("_" * 30)

    r2 = t.cell(1, 1).paragraphs[0]
    r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = r2.add_run("Authorized Signatory")
    rn.bold = True
    rn.font.size = Pt(10)

    r3 = t.cell(2, 1).paragraphs[0]
    r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3.add_run("Human Resources Department\nCraftsmen IT").font.size = Pt(9)

def add_footer_text(doc, ref_no, page="1 of 1"):
    add_hrule(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Craftsmen IT  |  IT Services Company in UAE & Saudi Arabia  |  Ref: {ref_no}  |  Page {page}")
    r.font.size = Pt(8)
    r.font.color.rgb = DARK_GRAY


# ════════════════════════════════════════════════════════════════════════════
# 1. JOB / EMPLOYMENT CERTIFICATE
# ════════════════════════════════════════════════════════════════════════════
def make_job_certificate():
    doc = new_doc()
    add_header_block(doc, "CMI/HR/EC/2026/034", "August 29, 2026")
    add_title(doc, "EMPLOYMENT CERTIFICATE")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run("To Whom It May Concern,")

    body = (
        "This is to certify that "
    )
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.space_after = Pt(8)
    p2.add_run("This is to certify that ")
    b = p2.add_run("Mr. Arafat Iqbal")
    b.bold = True
    p2.add_run(", bearing Employee ID ")
    b2 = p2.add_run("CMI0034")
    b2.bold = True
    p2.add_run(", has been employed with ")
    b3 = p2.add_run("Craftsmen IT")
    b3.bold = True
    p2.add_run(" as a ")
    b4 = p2.add_run("Trainee Software Engineer")
    b4.bold = True
    p2.add_run(" in the Software Engineering Department, effective from ")
    b5 = p2.add_run("19 February 2025")
    b5.bold = True
    p2.add_run(", and is currently on our active payroll as of the date of this certificate.")

    add_body_para(doc, "The details of his employment are as follows:")

    navy_table(doc,
        ["Particulars", "Details"],
        [
            ["Full Name",       "Arafat Iqbal"],
            ["Employee ID",     "CMI0034"],
            ["Designation",     "Trainee Software Engineer"],
            ["Department",      "Software Engineering"],
            ["Employment Type", "Full-Time"],
            ["Date of Joining", "19 February 2025"],
            ["Current Status",  "Active Employee"],
        ]
    )

    add_body_para(doc,
        "During his tenure with Craftsmen IT, Mr. Arafat Iqbal has demonstrated professionalism, "
        "dedication, and a strong commitment to his responsibilities. He has been a reliable member "
        "of our team and has contributed positively to the organization."
    )
    add_body_para(doc,
        "This certificate is being issued upon the request of the employee for whatever lawful purpose it may serve."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run("For further verification or inquiries, please contact our HR Department at ")
    b = p.add_run("contact@craftsmenit.com")
    b.bold = True
    p.add_run(" or call ")
    b2 = p.add_run("+966-59-787-9394")
    b2.bold = True
    p.add_run(".")

    add_signature_block(doc)
    add_footer_text(doc, "CMI/HR/EC/2026/034")
    doc.save(os.path.join(OUT, "1_Job_Certificate.docx"))
    print("✓ 1_Job_Certificate.docx")


# ════════════════════════════════════════════════════════════════════════════
# 2. SALARY CERTIFICATE
# ════════════════════════════════════════════════════════════════════════════
def make_salary_certificate():
    doc = new_doc()
    add_header_block(doc, "CMI/HR/SC/2026/034", "August 29, 2026")
    add_title(doc, "SALARY CERTIFICATE")

    doc.add_paragraph("To Whom It May Concern,").paragraph_format.space_after = Pt(8)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.space_after = Pt(8)
    p2.add_run("This is to certify that ")
    p2.add_run("Mr. Arafat Iqbal").bold = True
    p2 = doc.paragraphs[-1]  # re-grab
    p2.add_run(", Employee ID ")
    p2.add_run("CMI0034").bold = True
    p2 = doc.paragraphs[-1]
    p2.add_run(", is currently employed with ")
    p2.add_run("Craftsmen IT").bold = True
    p2 = doc.paragraphs[-1]
    p2.add_run(" in the capacity of ")
    p2.add_run("Trainee Software Engineer").bold = True
    p2 = doc.paragraphs[-1]
    p2.add_run(" since ")
    p2.add_run("19 February 2025").bold = True
    p2 = doc.paragraphs[-1]
    p2.add_run(". His monthly compensation package as of the date of this certificate is detailed below:")

    navy_table(doc,
        ["Salary Component", "Monthly Amount (SAR)"],
        [
            ["Basic Salary",              "SAR [BASIC_SALARY]"],
            ["Housing Allowance",         "SAR [HOUSING_ALLOWANCE]"],
            ["Transportation Allowance",  "SAR [TRANSPORT_ALLOWANCE]"],
            ["Other Allowances",          "SAR [OTHER_ALLOWANCES]"],
            ["TOTAL GROSS SALARY",        "SAR [TOTAL_GROSS]"],
        ]
    )

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(8)
    r = p3.add_run("Total Monthly Gross Salary (in words):  SAR [TOTAL_IN_WORDS] Only")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY

    add_body_para(doc,
        "His salary is credited directly to his designated bank account on a monthly basis. "
        "The above figures represent his gross compensation before any applicable statutory deductions."
    )
    add_body_para(doc,
        "This certificate is issued at the specific request of the employee and is valid for the purpose "
        "for which it has been requested."
    )

    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_after = Pt(8)
    r_note = p_note.add_run(
        "Note: This certificate is confidential and intended solely for the use of the addressee. "
        "Any reproduction or disclosure without written authorization from Craftsmen IT is prohibited."
    )
    r_note.italic = True
    r_note.font.size = Pt(9)
    r_note.font.color.rgb = DARK_GRAY

    add_signature_block(doc)
    add_footer_text(doc, "CMI/HR/SC/2026/034")
    doc.save(os.path.join(OUT, "2_Salary_Certificate.docx"))
    print("✓ 2_Salary_Certificate.docx")


# ════════════════════════════════════════════════════════════════════════════
# 3. NO OBJECTION CERTIFICATE
# ════════════════════════════════════════════════════════════════════════════
def make_noc():
    doc = new_doc()
    add_header_block(doc, "CMI/HR/NOC/2026/034", "August 29, 2026")
    add_title(doc, "NO OBJECTION CERTIFICATE (NOC)")

    doc.add_paragraph("To Whom It May Concern,").paragraph_format.space_after = Pt(8)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.space_after = Pt(8)
    p2.add_run("This is to certify that ")
    p2.add_run("Mr. Arafat Iqbal").bold = True
    p2 = doc.paragraphs[-1]
    p2.add_run(", Employee ID ")
    p2.add_run("CMI0034").bold = True
    p2 = doc.paragraphs[-1]
    p2.add_run(", is a bonafide employee of ")
    p2.add_run("Craftsmen IT").bold = True
    p2 = doc.paragraphs[-1]
    p2.add_run(", currently serving as a ")
    p2.add_run("Trainee Software Engineer").bold = True
    p2 = doc.paragraphs[-1]
    p2.add_run(" in the Software Engineering Department since ")
    p2.add_run("19 February 2025").bold = True
    p2 = doc.paragraphs[-1]
    p2.add_run(".")

    add_body_para(doc, "The details of the employee are as follows:")

    navy_table(doc,
        ["Particulars", "Details"],
        [
            ["Full Name",         "Arafat Iqbal"],
            ["Employee ID",       "CMI0034"],
            ["Designation",       "Trainee Software Engineer"],
            ["Department",        "Software Engineering"],
            ["Date of Joining",   "19 February 2025"],
            ["Employment Status", "Active / Full-Time"],
        ]
    )

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(8)
    p3.add_run("Craftsmen IT has ")
    b = p3.add_run("No Objection")
    b.bold = True
    p3.add_run(" to Mr. Arafat Iqbal applying for / availing the following:")

    for item in [
        "☐   Visa / Travel Purposes",
        "☐   Higher Education / Admission to an Academic Institution",
        "☐   Bank Loan / Financial Services",
        "☐   Any Other Official / Legal Purpose: ___________________________",
    ]:
        pi = doc.add_paragraph(item)
        pi.paragraph_format.space_after = Pt(4)
        pi.paragraph_format.left_indent = Inches(0.3)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_body_para(doc,
        "The company confirms that there is no legal, financial, or disciplinary case pending against "
        "the above-named employee. He is in good standing with the organization, and we have no objection "
        "to him undertaking the above-mentioned activity, provided that it does not interfere with his "
        "professional obligations to Craftsmen IT."
    )
    add_body_para(doc,
        "This certificate is issued in good faith at the request of the employee for lawful purposes only, "
        "and it does not constitute a guarantee, warranty, or indemnity of any kind by Craftsmen IT."
    )
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p4.add_run("For verification, please contact our HR Department at ")
    p4.add_run("contact@craftsmenit.com").bold = True
    p4 = doc.paragraphs[-1]
    p4.add_run(" or ")
    p4.add_run("+966-59-787-9394").bold = True
    p4 = doc.paragraphs[-1]
    p4.add_run(".")

    add_signature_block(doc)
    add_footer_text(doc, "CMI/HR/NOC/2026/034")
    doc.save(os.path.join(OUT, "3_NOC.docx"))
    print("✓ 3_NOC.docx")


# ════════════════════════════════════════════════════════════════════════════
# 4. APPOINTMENT LETTER
# ════════════════════════════════════════════════════════════════════════════
def make_appointment_letter():
    doc = new_doc()
    add_header_block(doc, "CMI/HR/APT/2025/034", "14 February 2025")
    add_title(doc, "APPOINTMENT LETTER")

    # Address block
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.add_run("Mr. Arafat Iqbal\n[Employee's Residential Address]\n[City, Country]")

    doc.add_paragraph("Dear Mr. Arafat Iqbal,").paragraph_format.space_after = Pt(8)

    add_body_para(doc,
        "We are pleased to offer you the position of Trainee Software Engineer at Craftsmen IT. "
        "Following the successful completion of your interviews and assessments, the management is "
        "confident that you will be a valuable addition to our team. The terms and conditions of "
        "your employment are outlined below:"
    )

    def section(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = NAVY
        r.font.size = Pt(11)

    section("1.  Employment Details")
    navy_table(doc,
        ["Particulars", "Details"],
        [
            ["Full Name",       "Arafat Iqbal"],
            ["Employee ID",     "CMI0034"],
            ["Designation",     "Trainee Software Engineer"],
            ["Department",      "Software Engineering"],
            ["Reporting To",    "[Manager Name / Team Lead]"],
            ["Date of Joining", "19 February 2025"],
            ["Employment Type", "Full-Time, Trainee"],
            ["Work Location",   "Craftsmen IT – Main Office"],
        ]
    )

    section("2.  Compensation & Benefits")
    navy_table(doc,
        ["Component", "Monthly Amount (SAR)"],
        [
            ["Basic Salary",             "SAR [BASIC_SALARY]"],
            ["Housing Allowance",        "SAR [HOUSING_ALLOWANCE]"],
            ["Transportation Allowance", "SAR [TRANSPORT_ALLOWANCE]"],
            ["Other Allowances",         "SAR [OTHER_ALLOWANCES]"],
            ["TOTAL GROSS MONTHLY",      "SAR [TOTAL_GROSS]"],
        ]
    )

    section("3.  Working Hours")
    add_body_para(doc,
        "Your standard working hours will be 9:00 AM to 6:00 PM, Monday through Friday, with a one-hour "
        "lunch break. Overtime may be required in exceptional business circumstances and will be compensated "
        "in accordance with the company policy."
    )

    section("4.  Probation Period")
    add_body_para(doc,
        "You will be on a probation period of Three (3) months from your date of joining, i.e., from "
        "19 February 2025 to 18 May 2025. During this period, either party may terminate the employment "
        "with 7 days' written notice. Upon successful completion of probation, you will be confirmed "
        "as a regular employee."
    )

    section("5.  Leave Entitlement")
    add_body_para(doc,
        "You will be entitled to 21 days of annual leave per year (pro-rated in the first year), "
        "10 days of public holidays, and sick leave as per applicable labor law. "
        "Leave is subject to approval by your line manager."
    )

    section("6.  Confidentiality & Non-Disclosure")
    add_body_para(doc,
        "As part of your employment, you will be required to sign a Confidentiality and Non-Disclosure "
        "Agreement. All information, data, processes, and intellectual property belonging to Craftsmen IT "
        "must be kept strictly confidential during and after your employment."
    )

    section("7.  Code of Conduct")
    add_body_para(doc,
        "You are expected to maintain the highest standards of professional conduct, integrity, and ethics "
        "while representing Craftsmen IT. Any violation of the company's code of conduct may result in "
        "disciplinary action, including termination of employment."
    )

    section("8.  Termination")
    add_body_para(doc,
        "After the probation period, either party may terminate this agreement by giving One (1) Month's "
        "written notice or by paying the equivalent salary in lieu of notice. The company reserves the right "
        "to terminate employment without notice in cases of gross misconduct or breach of company policy."
    )

    section("9.  Governing Law")
    add_body_para(doc,
        "This appointment letter and the terms of your employment shall be governed by and construed in "
        "accordance with the labor laws of the Kingdom of Saudi Arabia / UAE (as applicable to your work location)."
    )

    section("10. Acceptance")
    add_body_para(doc,
        "Please sign and return a copy of this letter by [Acceptance Deadline Date] to confirm your acceptance "
        "of the above terms and conditions. Failure to do so within the stipulated period will be considered as "
        "non-acceptance, and the offer shall stand withdrawn."
    )

    add_body_para(doc,
        "We look forward to welcoming you to the Craftsmen IT family. We are confident that this will be a "
        "mutually rewarding association and wish you a successful career with us."
    )

    doc.add_paragraph("Yours sincerely,").paragraph_format.space_after = Pt(4)
    add_signature_block(doc)
    add_hrule(doc)

    # Acceptance box
    p_acc = doc.add_paragraph()
    r_acc = p_acc.add_run("EMPLOYEE ACCEPTANCE DECLARATION")
    r_acc.bold = True
    r_acc.font.color.rgb = NAVY
    r_acc.font.size = Pt(11)

    add_body_para(doc,
        "I, Arafat Iqbal, have read, understood, and accept all the terms and conditions of employment "
        "as stated in this Appointment Letter."
    )

    for line in [
        "Signature:   _____________________________",
        "Date:          _____________________________",
        "Employee Name:   Arafat Iqbal",
        "Employee ID:       CMI0034",
    ]:
        doc.add_paragraph(line).paragraph_format.space_after = Pt(6)

    add_footer_text(doc, "CMI/HR/APT/2025/034")
    doc.save(os.path.join(OUT, "4_Appointment_Letter.docx"))
    print("✓ 4_Appointment_Letter.docx")


# ════════════════════════════════════════════════════════════════════════════
# 5. PAY SLIPS – LAST 6 MONTHS
# ════════════════════════════════════════════════════════════════════════════
MONTHS = [
    ("March 2026",   "01 Mar 2026", "31 Mar 2026"),
    ("April 2026",   "01 Apr 2026", "30 Apr 2026"),
    ("May 2026",     "01 May 2026", "31 May 2026"),
    ("June 2026",    "01 Jun 2026", "30 Jun 2026"),
    ("July 2026",    "01 Jul 2026", "31 Jul 2026"),
    ("August 2026",  "01 Aug 2026", "31 Aug 2026"),
]

def make_payslips():
    doc = new_doc()

    for idx, (month, start, end) in enumerate(MONTHS):
        if idx > 0:
            doc.add_page_break()

        # Mini header
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("CRAFTSMEN IT")
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = NAVY

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        p2.add_run("IT Services & Solutions  |  UAE & Saudi Arabia  |  contact@craftsmenit.com").font.size = Pt(9)

        add_hrule(doc)

        # Slip title + month
        t_title = doc.add_table(rows=1, cols=2)
        t_title.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell in t_title.rows[0].cells:
            for b in ('top','left','bottom','right'):
                set_cell_border(cell, **{b:'none'})
        lp = t_title.cell(0,0).paragraphs[0]
        lr = lp.add_run("PAY SLIP")
        lr.bold = True
        lr.font.size = Pt(14)
        lr.font.color.rgb = NAVY
        rp = t_title.cell(0,1).paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = rp.add_run(f"Month of {month}")
        rr.bold = True
        rr.font.size = Pt(11)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # Employee info table
        info_t = doc.add_table(rows=2, cols=4)
        info_t.style = 'Table Grid'
        info_t.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_data = [
            ("Employee Name", "Arafat Iqbal", "Employee ID", "CMI0034"),
            ("Designation", "Trainee Software Engineer", "Pay Period", f"{start} – {end}"),
        ]
        for ri, row_data in enumerate(info_data):
            for ci, val in enumerate(row_data):
                cell = info_t.rows[ri].cells[ci]
                set_cell_bg(cell, 'E8EDF5' if ci % 2 == 0 else 'FFFFFF')
                p = cell.paragraphs[0]
                r = p.add_run(val)
                r.font.size = Pt(10)
                if ci % 2 == 0:
                    r.bold = True
                    r.font.color.rgb = NAVY
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Earnings + Deductions side by side as one wide table
        earn_ded = doc.add_table(rows=6, cols=4)
        earn_ded.style = 'Table Grid'
        earn_ded.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = [("EARNINGS", "AMOUNT (SAR)", "DEDUCTIONS", "AMOUNT (SAR)")]
        e_rows = [
            ("Basic Salary", "SAR [BASIC]", "GOSI / Pension", "SAR [GOSI]"),
            ("Housing Allowance", "SAR [HOUSING]", "Absence Deduction", "SAR 0.00"),
            ("Transportation Allow.", "SAR [TRANSPORT]", "Loan / Advance", "SAR 0.00"),
            ("Other Allowances", "SAR [OTHER]", "Other Deductions", "SAR 0.00"),
            ("TOTAL EARNINGS", "SAR [TOTAL_EARN]", "TOTAL DEDUCTIONS", "SAR [TOTAL_DED]"),
        ]

        for ri, row_data in enumerate(headers + e_rows):
            is_hdr = ri == 0
            is_total = not is_hdr and row_data[0].startswith("TOTAL")
            for ci, val in enumerate(row_data):
                cell = earn_ded.rows[ri].cells[ci]
                if is_hdr or is_total:
                    set_cell_bg(cell, '1A3C6E')
                    r = cell.paragraphs[0].add_run(val)
                    r.bold = True
                    r.font.color.rgb = WHITE
                    r.font.size = Pt(10)
                else:
                    bg = 'F0F4FB' if ri % 2 == 0 else 'FFFFFF'
                    set_cell_bg(cell, bg)
                    cell.paragraphs[0].add_run(val).font.size = Pt(10)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Net Pay box
        p_net = doc.add_paragraph()
        p_net.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_net.paragraph_format.space_after = Pt(4)
        r_net = p_net.add_run(f"NET PAY FOR {month.upper()}:   SAR [NET_PAY]")
        r_net.bold = True
        r_net.font.size = Pt(13)
        r_net.font.color.rgb = NAVY

        # Amount in words
        p_w = doc.add_paragraph()
        p_w.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_w.paragraph_format.space_after = Pt(10)
        p_w.add_run("Amount in Words: SAR [NET_PAY_WORDS] Only   |   Payment Mode: Bank Transfer").font.size = Pt(9)

        # Signature row
        sig_t = doc.add_table(rows=2, cols=3)
        sig_t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell in sig_t.rows[0].cells:
            for b in ('top','left','bottom','right'):
                set_cell_border(cell, **{b:'none'})
            cell.paragraphs[0].add_run("_" * 22).font.size = Pt(10)
        labels = ["Employee Signature", "HR / Accounts", "Authorized Signatory"]
        for ci, lbl in enumerate(labels):
            cell = sig_t.rows[1].cells[ci]
            for b in ('top','left','bottom','right'):
                set_cell_border(cell, **{b:'none'})
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(lbl)
            r.bold = True
            r.font.size = Pt(9)

    doc.save(os.path.join(OUT, "5_Payslips_6_Months.docx"))
    print("✓ 5_Payslips_6_Months.docx")


# ════════════════════════════════════════════════════════════════════════════
# 6. SALARY STATEMENT – 6 MONTHS
# ════════════════════════════════════════════════════════════════════════════
def make_salary_statement():
    doc = new_doc()
    add_header_block(doc, "CMI/HR/SS/2026/034", "August 29, 2026")
    add_title(doc, "SALARY STATEMENT — LAST 6 MONTHS")

    # Employee info bar
    info_t = doc.add_table(rows=2, cols=3)
    info_t.style = 'Table Grid'
    info_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("Employee Name", "Arafat Iqbal", "Employee ID: CMI0034"),
        ("Designation", "Trainee Software Engineer", "Period: March – August 2026"),
    ]
    for ri, row_data in enumerate(info_data):
        for ci, val in enumerate(row_data):
            cell = info_t.rows[ri].cells[ci]
            set_cell_bg(cell, 'E8EDF5' if ci == 0 else 'FFFFFF')
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)
            if ci == 0:
                r.bold = True
                r.font.color.rgb = NAVY
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_body_para(doc,
        "This Salary Statement is issued upon the request of Mr. Arafat Iqbal (CMI0034) and reflects the "
        "salary disbursed for the period of March 2026 to August 2026. All amounts are in Saudi Arabian Riyal (SAR)."
    )

    navy_table(doc,
        ["Month", "Basic", "Housing", "Transport", "Other Allow.", "Gross", "Deductions", "Net Pay"],
        [
            ["March 2026",     "[BASIC]", "[HOUSING]", "[TRANSPORT]", "[OTHER]", "[GROSS]", "[DED]", "[NET]"],
            ["April 2026",     "[BASIC]", "[HOUSING]", "[TRANSPORT]", "[OTHER]", "[GROSS]", "[DED]", "[NET]"],
            ["May 2026",       "[BASIC]", "[HOUSING]", "[TRANSPORT]", "[OTHER]", "[GROSS]", "[DED]", "[NET]"],
            ["June 2026",      "[BASIC]", "[HOUSING]", "[TRANSPORT]", "[OTHER]", "[GROSS]", "[DED]", "[NET]"],
            ["July 2026",      "[BASIC]", "[HOUSING]", "[TRANSPORT]", "[OTHER]", "[GROSS]", "[DED]", "[NET]"],
            ["August 2026",    "[BASIC]", "[HOUSING]", "[TRANSPORT]", "[OTHER]", "[GROSS]", "[DED]", "[NET]"],
            ["TOTAL (6 Months)", "[T_BASIC]","[T_HOUSING]","[T_TRANSPORT]","[T_OTHER]","[T_GROSS]","[T_DED]","[T_NET]"],
        ]
    )

    p_total = doc.add_paragraph()
    p_total.paragraph_format.space_after = Pt(8)
    r_t = p_total.add_run("Total Net Amount Paid (6 Months):  SAR [TOTAL_NET_WORDS] Only")
    r_t.bold = True
    r_t.font.color.rgb = NAVY

    add_body_para(doc,
        "Payment for all months listed above was disbursed via bank transfer to the employee's designated bank account. "
        "All payments were made on or before the last working day of each respective month."
    )

    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_after = Pt(8)
    r_note = p_note.add_run(
        "Note: This statement is issued for official purposes only and is subject to verification. "
        "All amounts are gross of any applicable taxes unless stated otherwise."
    )
    r_note.italic = True
    r_note.font.size = Pt(9)
    r_note.font.color.rgb = DARK_GRAY

    add_signature_block(doc)
    add_footer_text(doc, "CMI/HR/SS/2026/034")
    doc.save(os.path.join(OUT, "6_Salary_Statement_6_Months.docx"))
    print("✓ 6_Salary_Statement_6_Months.docx")


# ════════════════════════════════════════════════════════════════════════════
# 7. OFFICE ID CARD + VISITING CARD  (text-based, print on card stock)
# ════════════════════════════════════════════════════════════════════════════
def make_id_card():
    doc = new_doc()

    p = doc.add_paragraph()
    r = p.add_run("OFFICE ID CARD & VISITING / BUSINESS CARD")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Craftsmen IT  —  Employee: Arafat Iqbal (CMI0034)").paragraph_format.space_after = Pt(2)
    add_hrule(doc)

    # ── ID CARD ──────────────────────────────────────────────────────────
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(4)
    p1.paragraph_format.space_before = Pt(8)
    p1.add_run("EMPLOYEE ID CARD  (Print at 85.6 × 54 mm — Standard CR80 Card)").bold = True
    p1.runs[0].font.color.rgb = NAVY

    id_t = doc.add_table(rows=1, cols=2)
    id_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    id_t.style = 'Table Grid'

    # Front cell
    front = id_t.cell(0, 0)
    set_cell_bg(front, '1A3C6E')
    fp = front.paragraphs[0]
    fr = fp.add_run(
        "CRAFTSMEN IT\n"
        "━━━━━━━━━━━━━━━━━━━━━━\n"
        "\n"
        "   [ PHOTO ]\n"
        "\n"
        "Name:       Arafat Iqbal\n"
        "Title:         Trainee Software Engineer\n"
        "Emp ID:    CMI0034\n"
        "Dept:         Software Engineering\n"
        "Joined:      19 Feb 2025\n"
        "Blood Grp: [BLOOD_GROUP]\n"
        "\n"
        "━━━━━━━━━━━━━━━━━━━━━━\n"
        "contact@craftsmenit.com"
    )
    fr.font.color.rgb = WHITE
    fr.font.size = Pt(10)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Back cell
    back = id_t.cell(0, 1)
    set_cell_bg(back, 'FFFFFF')
    bp = back.paragraphs[0]
    br = bp.add_run(
        "CRAFTSMEN IT\n"
        "Employee Identification\n"
        "━━━━━━━━━━━━━━━━━━━━━━\n"
        "\n"
        "Name:    Arafat Iqbal\n"
        "ID:          CMI0034\n"
        "DOJ:       19 February 2025\n"
        "Email:    arafat@craftsmenit.com\n"
        "\n"
        "Emergency Contact:\n"
        "[CONTACT_NAME] — [PHONE]\n"
        "\n"
        "━━━━━━━━━━━━━━━━━━━━━━\n"
        "If found, return to Craftsmen IT\n"
        "www.craftsmenit.com"
    )
    br.font.size = Pt(10)
    br.font.color.rgb = NAVY
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Label row
    label_t = doc.add_table(rows=1, cols=2)
    label_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, lbl in enumerate(["FRONT — ID Card", "BACK — ID Card"]):
        cell = label_t.cell(0, ci)
        for b in ('top','left','bottom','right'):
            set_cell_border(cell, **{b:'none'})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(lbl)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = DARK_GRAY

    add_hrule(doc)

    # ── VISITING / BUSINESS CARD ──────────────────────────────────────────
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(8)
    p2.paragraph_format.space_after = Pt(4)
    p2.add_run("VISITING / BUSINESS CARD  (Print at 90 × 55 mm — Standard Business Card)").bold = True
    p2.runs[0].font.color.rgb = NAVY

    vc_t = doc.add_table(rows=1, cols=2)
    vc_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    vc_t.style = 'Table Grid'

    # Front
    vf = vc_t.cell(0, 0)
    set_cell_bg(vf, '1A3C6E')
    vfp = vf.paragraphs[0]
    vfr = vfp.add_run(
        "CRAFTSMEN IT\n"
        "━━━━━━━━━━━━━━━━━━━━━━━━━\n"
        "\n"
        "Arafat Iqbal\n"
        "Trainee Software Engineer\n"
        "\n"
        "✉  arafat@craftsmenit.com\n"
        "☎  +966-59-787-9394\n"
        "🌐  www.craftsmenit.com"
    )
    vfr.font.color.rgb = WHITE
    vfr.font.size = Pt(10)
    vfp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Back
    vb = vc_t.cell(0, 1)
    set_cell_bg(vb, 'FFFFFF')
    vbp = vb.paragraphs[0]
    vbr = vbp.add_run(
        "\n"
        "CRAFTSMEN IT\n"
        "IT Services & Solutions\n"
        "\n"
        "www.craftsmenit.com\n"
        "\n"
        '"The right technology,\n implemented properly."'
    )
    vbr.font.size = Pt(11)
    vbr.font.color.rgb = NAVY
    vbp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    label_t2 = doc.add_table(rows=1, cols=2)
    label_t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, lbl in enumerate(["FRONT — Business Card", "BACK — Business Card"]):
        cell = label_t2.cell(0, ci)
        for b in ('top','left','bottom','right'):
            set_cell_border(cell, **{b:'none'})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(lbl)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = DARK_GRAY

    add_hrule(doc)
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_r = p_inst.add_run(
        "Replace all [PLACEHOLDER] values with actual data before printing.\n"
        "Print at 100% scale. Use card stock for ID and business cards."
    )
    inst_r.italic = True
    inst_r.font.size = Pt(9)
    inst_r.font.color.rgb = DARK_GRAY

    doc.save(os.path.join(OUT, "7_Office_ID_and_Visiting_Card.docx"))
    print("✓ 7_Office_ID_and_Visiting_Card.docx")


# ─── run all ────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    make_job_certificate()
    make_salary_certificate()
    make_noc()
    make_appointment_letter()
    make_payslips()
    make_salary_statement()
    make_id_card()
    print("\nAll documents saved to:", OUT)
